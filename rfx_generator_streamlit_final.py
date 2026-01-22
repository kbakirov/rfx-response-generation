import streamlit as st
import anthropic
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
from datetime import datetime
import re

# Page config
st.set_page_config(
    page_title="RFX Response Generator",
    page_icon="📄",
    layout="wide"
)

# Initialize session state
if 'step' not in st.session_state:
    st.session_state.step = 1
if 'rfx_data' not in st.session_state:
    st.session_state.rfx_data = {}
if 'generated_sections' not in st.session_state:
    st.session_state.generated_sections = {}

# Sections configuration
SECTIONS = [
    {'key': 'executive_summary', 'label': 'Executive summary', 'icon': '📋'},
    {'key': 'implementation_architecture', 'label': 'Implementation architecture', 'icon': '🏗️'},
    {'key': 'team_members', 'label': 'Team composition', 'icon': '👥'},
    {'key': 'cost_estimate', 'label': 'Cost estimate', 'icon': '💰'},
    {'key': 'execution_plan', 'label': 'Execution plan', 'icon': '📅'},
    {'key': 'timeline', 'label': 'Timeline and Milestones', 'icon': '⏱️'},
    {'key': 'risks', 'label': 'Risk assessment', 'icon': '⚠️'},
    {'key': 'assumptions', 'label': 'Assumptions and Dependencies', 'icon': '📌'},
    {'key': 'deliverables', 'label': 'Deliverables', 'icon': '✅'},
    {'key': 'conclusion', 'label': 'Conclusion', 'icon': '🎯'}
]

def clean_markdown_text(text):
    """Remove markdown formatting from text"""
    # Remove bold/italic markers
    text = re.sub(r'\*\*\*(.+?)\*\*\*', r'\1', text)  # Bold italic
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)  # Bold
    text = re.sub(r'\*(.+?)\*', r'\1', text)  # Italic
    text = re.sub(r'__(.+?)__', r'\1', text)  # Bold underscore
    text = re.sub(r'_(.+?)_', r'\1', text)  # Italic underscore
    
    # Remove headers
    text = re.sub(r'^#{1,6}\s+', '', text, flags=re.MULTILINE)
    
    # Remove extra asterisks and markers
    text = re.sub(r'\*+', '', text)
    
    return text.strip()

def parse_markdown_table(text):
    """Extract tables from markdown text"""
    tables = []
    lines = text.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        # Check if line looks like a table row
        if '|' in line and line.count('|') >= 2:
            table_rows = []
            # Collect all consecutive table rows
            while i < len(lines) and '|' in lines[i]:
                row = lines[i].strip()
                if row and not re.match(r'^\|[\s\-:]+\|', row):  # Skip separator rows
                    cells = [cell.strip() for cell in row.split('|') if cell.strip()]
                    if cells:
                        table_rows.append(cells)
                i += 1
            if table_rows:
                tables.append(table_rows)
            continue
        i += 1
    
    return tables

def add_table_to_doc(doc, table_data):
    """Add a properly formatted table to the document"""
    if not table_data:
        return
    
    # Create table
    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
    table.style = 'Light Grid Accent 1'
    
    # Populate table
    for i, row_data in enumerate(table_data):
        row = table.rows[i]
        for j, cell_data in enumerate(row_data):
            cell = row.cells[j]
            cell.text = clean_markdown_text(str(cell_data))
            
            # Header row formatting
            if i == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(11)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Add shading to header
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), 'D5E8F0')
                cell._element.get_or_add_tcPr().append(shading)
            else:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)

def add_formatted_content(doc, content):
    """Add content to document with proper formatting, handling markdown"""
    # First, extract and process tables
    tables = parse_markdown_table(content)
    
    # Remove table content from text
    text_without_tables = content
    for table in tables:
        # Find and remove the table text
        table_pattern = r'\|[^\n]+\|(\n\|[^\n]+\|)+'
        text_without_tables = re.sub(table_pattern, '[TABLE_PLACEHOLDER]', text_without_tables, count=1)
    
    # Split content into paragraphs
    paragraphs = text_without_tables.split('\n\n')
    
    table_index = 0
    for para_text in paragraphs:
        if not para_text.strip():
            continue
            
        # Check if this is a table placeholder
        if '[TABLE_PLACEHOLDER]' in para_text:
            if table_index < len(tables):
                add_table_to_doc(doc, tables[table_index])
                doc.add_paragraph()  # Add space after table
                table_index += 1
            continue
        
        # Clean markdown formatting
        para_text = clean_markdown_text(para_text)
        
        # Check if it's a heading (starts with original ### or ##)
        if para_text.startswith('###'):
            para_text = para_text.replace('###', '').strip()
            p = doc.add_paragraph(para_text)
            p.style = 'Heading 3'
        elif para_text.startswith('##'):
            para_text = para_text.replace('##', '').strip()
            p = doc.add_paragraph(para_text)
            p.style = 'Heading 2'
        else:
            # Regular paragraph with LEFT alignment (not justified - fixes spacing issues)
            p = doc.add_paragraph(para_text)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Changed from JUSTIFY to LEFT
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            p.paragraph_format.line_spacing = 1.15
            
            # Set font
            for run in p.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(11)

def get_section_prompt(section_key):
    """Get specific prompts for each section"""
    prompts = {
        'executive_summary': 'Provide a concise overview of the proposed solution, highlighting key benefits and approach. Keep it to 2-3 paragraphs.',
        'implementation_architecture': 'Detail the technical architecture including: system components, technology stack, integration points, infrastructure requirements, scalability considerations, and security measures. Be specific about technologies and frameworks.',
        'team_members': 'Propose a realistic team structure including: roles (e.g., Solution Architect, Senior Developer, QA Engineer), number of resources per role, allocation percentages, and key responsibilities. Include realistic experience levels.',
        'cost_estimate': 'Provide a detailed cost breakdown including: labor costs by role, infrastructure costs, software licenses, third-party services, contingency (10-15%), and total project cost. Use realistic hourly rates and estimates.',
        'execution_plan': 'Create a phased execution plan with: Phase descriptions, activities per phase, entry/exit criteria, and parallel work streams if applicable.',
        'timeline': 'Create a realistic project timeline with: major milestones, durations for each phase, dependencies, and critical path items. Use weeks/months as appropriate.',
        'risks': 'Identify 5-7 realistic project risks with: risk description, probability (Low/Medium/High), impact (Low/Medium/High), and mitigation strategy for each.',
        'assumptions': 'List 5-8 key assumptions and dependencies including: client responsibilities, data availability, infrastructure access, resource availability, and third-party dependencies.',
        'deliverables': 'List specific project deliverables organized by phase including: documentation, code deliverables, environments, training materials, and handover items.',
        'conclusion': 'Write a strong closing statement that: reinforces value proposition, expresses commitment, includes call-to-action, and contact information.'
    }
    return prompts.get(section_key, 'Generate comprehensive content for this section.')

def generate_section(client, section_key, section_label, rfx_data):
    """Generate a section using Claude API"""
    prompt = f"""You are an expert IT consultant creating an RFX response for a client.

Client Details:
- Client Name: {rfx_data.get('client_name', 'N/A')}
- Project Title: {rfx_data.get('project_title', 'N/A')}
- Project Type: {rfx_data.get('project_type', 'N/A')}
- Budget Range: {rfx_data.get('budget', 'N/A')}
- Timeline: {rfx_data.get('timeline', 'N/A')}

Project Description:
{rfx_data.get('description', 'N/A')}

Requirements:
{rfx_data.get('requirements', 'N/A')}

Constraints:
{rfx_data.get('constraints', 'N/A')}

Generate a detailed, professional, and realistic "{section_label}" section for this RFX response.

{get_section_prompt(section_key)}

Make it specific to this project, realistic, and professional. Use concrete numbers, specific technologies, and realistic timelines. Format with clear headings and structure."""

    try:
        message = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=1500,  # Reduced from 2000 for faster generation
            messages=[{"role": "user", "content": prompt}]
        )
        return message.content[0].text
    except Exception as e:
        return f"Error generating section: {str(e)}"

def generate_all_sections(client, rfx_data):
    """Generate all sections concurrently for faster performance"""
    from concurrent.futures import ThreadPoolExecutor, as_completed
    import time
    
    sections_content = {}
    progress_bar = st.progress(0)
    status_text = st.empty()
    
    # Create a thread pool for parallel generation
    with ThreadPoolExecutor(max_workers=5) as executor:
        # Submit all section generation tasks
        future_to_section = {
            executor.submit(generate_section, client, section['key'], section['label'], rfx_data): section
            for section in SECTIONS
        }
        
        completed = 0
        total = len(SECTIONS)
        
        # Process completed tasks as they finish
        for future in as_completed(future_to_section):
            section = future_to_section[future]
            try:
                content = future.result()
                sections_content[section['key']] = content
                completed += 1
                progress_bar.progress(completed / total)
                status_text.text(f"Generated {completed}/{total} sections... ({section['label']})")
            except Exception as e:
                st.error(f"Error generating {section['label']}: {str(e)}")
                sections_content[section['key']] = f"Error: {str(e)}"
                completed += 1
                progress_bar.progress(completed / total)
    
    status_text.text("✅ All sections generated!")
    return sections_content

def create_docx_document(rfx_data, sections_content):
    """Create a professional DOCX document with proper formatting"""
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Title page
    title_para = doc.add_paragraph()
    title_run = title_para.add_run('RFX RESPONSE DOCUMENT')
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 51, 102)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.space_after = Pt(24)
    
    # Add some space
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Client Information Box
    info_para = doc.add_paragraph()
    info_para.paragraph_format.space_before = Pt(12)
    info_para.paragraph_format.space_after = Pt(12)
    
    # Create info table
    info_table = doc.add_table(rows=4, cols=2)
    info_table.style = 'Light Grid Accent 1'
    
    info_data = [
        ('Client:', rfx_data.get('client_name', 'N/A')),
        ('Project:', rfx_data.get('project_title', 'N/A')),
        ('Type:', rfx_data.get('project_type', 'N/A')),
        ('Date:', datetime.now().strftime('%B %d, %Y'))
    ]
    
    for i, (label, value) in enumerate(info_data):
        row = info_table.rows[i]
        
        # Label cell
        label_cell = row.cells[0]
        label_cell.text = label
        for paragraph in label_cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(11)
        
        # Add shading to label cell
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'E7E6E6')
        label_cell._element.get_or_add_tcPr().append(shading)
        
        # Value cell
        value_cell = row.cells[1]
        value_cell.text = value
        for paragraph in value_cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(11)
    
    doc.add_page_break()
    
    # Add all sections with proper formatting
    for section in SECTIONS:
        # Section heading
        heading = doc.add_heading(section['label'], 1)
        heading.paragraph_format.space_before = Pt(18)
        heading.paragraph_format.space_after = Pt(12)
        
        # Format heading
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0, 51, 102)
            run.font.size = Pt(16)
        
        # Section content with proper formatting
        content = sections_content.get(section['key'], 'Content not generated')
        add_formatted_content(doc, content)
        
        # Add space after section
        doc.add_paragraph()
    
    # Footer
    doc.add_page_break()
    
    footer_heading = doc.add_heading('Contact Information', 1)
    footer_para = doc.add_paragraph()
    footer_para.add_run('For questions or clarifications regarding this proposal, please contact:\n\n')
    footer_para.add_run('Itransition\n').bold = True
    footer_para.add_run('Email: sales@itransition.com\n')
    footer_para.add_run('Phone: +1 (555) 123-4567\n')
    footer_para.add_run('Website: www.itransition.com')
    
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.paragraph_format.space_before = Pt(24)
    
    # Save document
    filename = f"RFX_{rfx_data.get('client_name', 'Client').replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(filename)
    return filename

def check_libreoffice_installed():
    """Check if LibreOffice is installed and available"""
    try:
        import subprocess
        result = subprocess.run(['soffice', '--version'], 
                              capture_output=True, 
                              text=True, 
                              timeout=5)
        return result.returncode == 0
    except (FileNotFoundError, subprocess.TimeoutExpired):
        return False
    except Exception:
        return False

def convert_docx_to_pdf(docx_filename):
    """Convert DOCX to PDF using LibreOffice"""
    try:
        import subprocess
        
        # Get the directory and base filename
        directory = os.path.dirname(docx_filename) or '.'
        
        # Use LibreOffice to convert DOCX to PDF
        result = subprocess.run([
            'soffice',
            '--headless',
            '--convert-to',
            'pdf',
            '--outdir',
            directory,
            docx_filename
        ], capture_output=True, text=True, timeout=30)
        
        if result.returncode == 0:
            # Generate PDF filename
            pdf_filename = docx_filename.replace('.docx', '.pdf')
            if os.path.exists(pdf_filename):
                return pdf_filename
        
        return None
    except Exception as e:
        return None

# Check LibreOffice availability at startup
LIBREOFFICE_AVAILABLE = check_libreoffice_installed()

# Header
st.title("🚀 RFX response generator")
st.markdown("### AI-Powered proposal creation for Itransition projects")
st.divider()

# Progress indicator
col1, col2, col3 = st.columns(3)
steps = [
    ("1️⃣ Project Info", col1),
    ("2️⃣ Generate", col2),
    ("3️⃣ Review & Export", col3)
]

for idx, (label, col) in enumerate(steps, 1):
    with col:
        if st.session_state.step == idx:
            st.markdown(f"**:blue[{label}]**")
        elif st.session_state.step > idx:
            st.markdown(f"**:green[{label} ✓]**")
        else:
            st.markdown(f":gray[{label}]")

st.divider()

# Step 1: Project Information
if st.session_state.step == 1:
    st.header("📋 Step 1: Project Information")
    st.markdown("Provide details about the client and project requirements")
    
    col1, col2 = st.columns(2)
    
    with col1:
        client_name = st.text_input("Client Name *", value=st.session_state.rfx_data.get('client_name', ''))
        project_title = st.text_input("Project Title *", value=st.session_state.rfx_data.get('project_title', ''))
        project_type = st.selectbox(
            "Project Type *",
            ['', 'Cloud Migration', 'Application Modernization', 'Data Center Transformation',
             'Digital Transformation', 'System Integration', 'Custom Development',
             'DevOps Implementation', 'Security Enhancement'],
            index=0 if not st.session_state.rfx_data.get('project_type') else 
                  ['', 'Cloud Migration', 'Application Modernization', 'Data Center Transformation',
                   'Digital Transformation', 'System Integration', 'Custom Development',
                   'DevOps Implementation', 'Security Enhancement'].index(st.session_state.rfx_data.get('project_type'))
        )
    
    with col2:
        budget = st.text_input("Budget Range", value=st.session_state.rfx_data.get('budget', ''), 
                              placeholder="e.g., $200K - $500K")
        timeline = st.text_input("Expected Timeline", value=st.session_state.rfx_data.get('timeline', ''),
                                placeholder="e.g., 6 months")
    
    description = st.text_area(
        "Project Description *",
        value=st.session_state.rfx_data.get('description', ''),
        height=150,
        placeholder="Describe the project, current state, desired outcomes, and business objectives..."
    )
    
    requirements = st.text_area(
        "Key Requirements",
        value=st.session_state.rfx_data.get('requirements', ''),
        height=100,
        placeholder="List specific technical requirements, compliance needs, integration points..."
    )
    
    constraints = st.text_area(
        "Constraints & Limitations",
        value=st.session_state.rfx_data.get('constraints', ''),
        height=100,
        placeholder="List any constraints: existing systems, technology preferences, regulatory requirements..."
    )
    
    if st.button("Continue to Generation →", type="primary"):
        if client_name and project_title and project_type and description:
            st.session_state.rfx_data = {
                'client_name': client_name,
                'project_title': project_title,
                'project_type': project_type,
                'budget': budget,
                'timeline': timeline,
                'description': description,
                'requirements': requirements,
                'constraints': constraints
            }
            st.session_state.step = 2
            st.rerun()
        else:
            st.error("Please fill in all required fields (*)")

# Step 2: Generate Sections
elif st.session_state.step == 2:
    st.header("✨ Step 2: Generate RFX Response")
    st.markdown("AI will generate comprehensive sections for your proposal")
    
    # Display project summary
    with st.expander("📊 Project Summary", expanded=True):
        col1, col2 = st.columns(2)
        with col1:
            st.markdown(f"**Client:** {st.session_state.rfx_data.get('client_name')}")
            st.markdown(f"**Project:** {st.session_state.rfx_data.get('project_title')}")
        with col2:
            st.markdown(f"**Type:** {st.session_state.rfx_data.get('project_type')}")
            st.markdown(f"**Timeline:** {st.session_state.rfx_data.get('timeline', 'Not specified')}")
    
    # Show sections to be generated
    st.subheader("Sections to Generate:")
    cols = st.columns(2)
    for idx, section in enumerate(SECTIONS):
        with cols[idx % 2]:
            st.markdown(f"{section['icon']} {section['label']}")
    
    col1, col2 = st.columns([1, 3])
    with col1:
        if st.button("← Back"):
            st.session_state.step = 1
            st.rerun()
    
    with col2:
        # API Key input
        api_key = st.text_input("Anthropic API Key (optional - app handles this)", type="password",
                               help="In production, this would be handled securely on the backend")
        
        api_key = "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"
        if st.button("🎯 Generate RFX Response", type="primary"):
            if not api_key:
                st.info("ℹ️ In production deployment, the API key would be securely stored on the server. For demo purposes, please provide your Anthropic API key.")
            else:
                with st.spinner("Generating sections in parallel... This will take 30-60 seconds..."):
                    try:
                        client = anthropic.Anthropic(api_key=api_key)
                        sections_content = generate_all_sections(client, st.session_state.rfx_data)
                        st.session_state.generated_sections = sections_content
                        st.session_state.step = 3
                        st.success("✅ All sections generated successfully!")
                        st.rerun()
                    except Exception as e:
                        st.error(f"Error: {str(e)}")

# Step 3: Review & Export
elif st.session_state.step == 3:
    st.header("📝 Step 3: Review & Export")
    st.markdown("Review generated content, make edits, and export your RFX document")
    
    # Edit sections
    st.subheader("📋 Edit Sections")
    for section in SECTIONS:
        with st.expander(f"{section['icon']} {section['label']}", expanded=False):
            content = st.text_area(
                "Content",
                value=st.session_state.generated_sections.get(section['key'], ''),
                height=300,
                key=f"edit_{section['key']}"
            )
            st.session_state.generated_sections[section['key']] = content
    
    st.divider()
    
    # Export options
    st.subheader("📥 Export Document")
    
    # Show LibreOffice status
    if not LIBREOFFICE_AVAILABLE:
        st.info("ℹ️ **PDF Export Not Available**: LibreOffice is not installed. You can export as DOCX and convert to PDF using Microsoft Word or an online converter like [Smallpdf](https://smallpdf.com/docx-to-pdf) or [PDF.io](https://pdf.io/docx-to-pdf).")
    
    # Adjust columns based on LibreOffice availability
    if LIBREOFFICE_AVAILABLE:
        col1, col2, col3 = st.columns(3)
    else:
        col1, col2 = st.columns(2)
    
    with col1:
        if st.button("📄 Export as DOCX", type="primary", use_container_width=True):
            try:
                with st.spinner("Generating DOCX document..."):
                    filename = create_docx_document(st.session_state.rfx_data, st.session_state.generated_sections)
                    
                    with open(filename, 'rb') as f:
                        st.download_button(
                            label="⬇️ Download DOCX",
                            data=f,
                            file_name=filename,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )
                    st.success("✅ DOCX document generated successfully!")
            except Exception as e:
                st.error(f"Error generating DOCX: {str(e)}")
    
    # Only show PDF options if LibreOffice is available
    if LIBREOFFICE_AVAILABLE:
        with col2:
            if st.button("📕 Export as PDF", type="primary", use_container_width=True):
                try:
                    with st.spinner("Generating PDF document..."):
                        # First create DOCX
                        docx_filename = create_docx_document(st.session_state.rfx_data, st.session_state.generated_sections)
                        
                        # Convert to PDF
                        pdf_filename = convert_docx_to_pdf(docx_filename)
                        
                        if pdf_filename and os.path.exists(pdf_filename):
                            with open(pdf_filename, 'rb') as f:
                                st.download_button(
                                    label="⬇️ Download PDF",
                                    data=f,
                                    file_name=os.path.basename(pdf_filename),
                                    mime="application/pdf",
                                    use_container_width=True
                                )
                            st.success("✅ PDF document generated successfully!")
                        else:
                            st.error("⚠️ PDF conversion failed. Please download DOCX and convert manually.")
                except Exception as e:
                    st.error(f"Error generating PDF: {str(e)}")
                    st.info("💡 Tip: Download the DOCX and convert to PDF using Microsoft Word or an online converter.")
        
        with col3:
            if st.button("📦 Export Both Formats", type="secondary", use_container_width=True):
                try:
                    with st.spinner("Generating documents..."):
                        # Create DOCX
                        docx_filename = create_docx_document(st.session_state.rfx_data, st.session_state.generated_sections)
                        
                        col_a, col_b = st.columns(2)
                        
                        # DOCX download
                        with col_a:
                            with open(docx_filename, 'rb') as f:
                                st.download_button(
                                    label="⬇️ Download DOCX",
                                    data=f,
                                    file_name=docx_filename,
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    use_container_width=True
                                )
                        
                        # Try PDF conversion
                        pdf_filename = convert_docx_to_pdf(docx_filename)
                        if pdf_filename and os.path.exists(pdf_filename):
                            with col_b:
                                with open(pdf_filename, 'rb') as f:
                                    st.download_button(
                                        label="⬇️ Download PDF",
                                        data=f,
                                        file_name=os.path.basename(pdf_filename),
                                        mime="application/pdf",
                                        use_container_width=True
                                    )
                        
                        st.success("✅ Documents generated successfully!")
                except Exception as e:
                    st.error(f"Error generating documents: {str(e)}")
    else:
        # Show installation instructions when LibreOffice is not available
        with col2:
            with st.expander("📋 How to Install LibreOffice for PDF Export"):
                st.markdown("""
                **Windows:**
                1. Download from [libreoffice.org](https://www.libreoffice.org/download/download/)
                2. Run the installer
                3. Restart the application
                
                **macOS:**
                ```bash
                brew install libreoffice
                ```
                
                **Linux/Ubuntu:**
                ```bash
                sudo apt-get install libreoffice
                ```
                
                **Alternative:** Export DOCX and use online converters:
                - [Smallpdf.com](https://smallpdf.com/docx-to-pdf)
                - [PDF.io](https://pdf.io/docx-to-pdf)
                - Microsoft Word (File → Save As → PDF)
                """)
    
    st.divider()
    
    # Back button
    if st.button("← Back to Generate"):
        st.session_state.step = 2
        st.rerun()

# Footer
st.divider()
st.markdown("""
<div style='text-align: center; color: gray; padding: 20px;'>
    <p>Powered by Claude AI • Built for Itransition Excellence</p>
    <p style='font-size: 0.8em;'>© 2026 RFX Response Generator</p>
</div>
""", unsafe_allow_html=True)
